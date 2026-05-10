"""
Rewrite the two review docx files (internal + external) so they describe
Nesterenko V.V.'s thesis on vector embedding models for semantic search.

Critical layout constraints (each review = 1 page):
  - Internal: 14pt, 4 body paragraphs + 1 conclusion. We MUST fit our content
    into the existing 4 body slots — no new paragraphs are inserted, so the
    page-break behaviour from the original template is preserved.
  - External: 12pt, 5 body paragraphs + 1 conclusion. 1:1 mapping.

We never create new w:p elements — that's how earlier runs broke the
indentation and squeezed the last paragraph against the conclusion. Only
set_text() into existing paragraphs.
"""
import io
import shutil
import sys
import pathlib

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

ROOT = pathlib.Path("D:/repos/kb-semantic-search-benchmark")
THESIS = ROOT / "thesis"
BACKUPS = THESIS / "_backups"

INT_PATH = THESIS / "2026_М_ПІ_ІПЗм-24-1_Нестеренко_В_В_рецензія_внутрішня.docx"
EXT_PATH = THESIS / "2026_М_ПІ_ІПЗм-24-1_Нестеренко_В_В_рецензія_зовнішня.docx"

GROUP_LINE = (
    "групи ІПЗм-24-1 Нестеренка Віталія Вячеславовича "
    "спеціальність – 121- Інженерія програмного забезпечення"
)
ONP_LINE = "освітньо-наукова програма «Інженерія програмного забезпечення»"
THEME = (
    "«Дослідження моделей векторних ембеддінгів для семантичного пошуку "
    "текстових даних у корпоративних базах знань»"
)
GRADE_PLACEHOLDER = "«____»"

# ---------------------------------------------------------------------------
# INTERNAL — 4 body paragraphs (slots 9, 10, 11, 12), then conclusion @ 13.
# Compress: actuality / novelty / (methodology+practical) / drawbacks.
# ---------------------------------------------------------------------------
INTERNAL_BODY = [
    # 9 — actuality
    "Дослідження є актуальним у контексті стрімкого зростання обсягів "
    "корпоративних знань у текстовій формі та обмеженості традиційного "
    "лексичного пошуку, який погано справляється з синонімією, "
    "перефразуванням і мультимовним контентом. Здобувач обґрунтував "
    "потребу в застосуванні сучасних моделей векторних ембеддінгів для "
    "побудови систем семантичного пошуку, орієнтованих на корпоративні "
    "бази знань, у тому числі україномовні.",

    # 10 — scientific novelty
    "Наукова новизна полягає у комплексному порівняльному дослідженні "
    "чотирьох сучасних моделей векторних ембеддінгів — multilingual-E5-base, "
    "BGE-M3, Qwen3-Embedding-0.6B та nomic-embed-text-v1.5 — у поєднанні з "
    "лексичною baseline-моделлю BM25, проведеному на трьох різнопрофільних "
    "доменах знань (технічна документація, юридичні тексти, медичні "
    "матеріали). Достовірність результатів підтверджується застосуванням "
    "бутстреп-оцінки 95% довірчих інтервалів (n = 2000) для метрик якості "
    "ранжування.",

    # 11 — methodology + practical (merged)
    "Методологічна частина виконана коректно: для індексації використано "
    "бібліотеку FAISS з точним індексом IndexFlatIP та L2-нормалізованими "
    "векторами (еквівалент косинусної міри), якість оцінено за стандартними "
    "IR-метриками nDCG@10, MRR@10, Recall@10, Precision@10, для вибору "
    "моделі застосовано Парето-фронт і лінійну адитивну модель MCDA. "
    "Дослідження супроводжується реалізацією веб-застосунку на базі Flask, "
    "що охоплює повний цикл — чанкування, побудову індексів, семантичний "
    "пошук, обчислення метрик та аналіз результатів за окремими запитами.",

    # 12 — drawbacks
    "Недоліком роботи можна вважати обмежену кількість запитів і "
    "релевантних документів у бенчмарках за окремими доменами, а також "
    "відсутність експериментів із дотренуванням (fine-tuning) моделей на "
    "доменних даних, що могло б додатково підвищити якість пошуку. Попри "
    "зазначене, робота є цілісною, завершеною та має наукову і практичну "
    "цінність.",
]

INTERNAL_CONCLUSION = (
    f"Кваліфікаційна робота магістра, здобувача групи ІПЗм-24-1 "
    f"Нестеренка Віталія Вячеславовича відповідає вимогам до кваліфікаційних "
    f"робіт і заслуговує оцінки {GRADE_PLACEHOLDER}. Кваліфікаційну роботу "
    f"здобувача групи ІПЗм-24-1 Нестеренка В.В. можна представити до захисту "
    f"в ЕК за спеціальністю 121 – «Інженерія програмного забезпечення», "
    f"освітньо-наукова програма «Інженерія програмного забезпечення»."
)

INTERNAL_REVIEWER = "к.т.н., доцент каф. Інформатики ХНУРЕ\tІрина ВЕЧІРСЬКА"

# ---------------------------------------------------------------------------
# EXTERNAL — 5 body paragraphs (slots 13..17), then conclusion @ 18.
# ---------------------------------------------------------------------------
EXTERNAL_BODY = [
    # 13 — volume
    "Обсяг роботи – достатній. Розділи добре структуровані, змістовні. "
    "Надані усі необхідні додатки, що допомагають повною мірою оцінити "
    "виконану роботу.",

    # 14 — fit + applicability
    "З урахуванням складності, робота відповідає вимогам до кваліфікаційної "
    "роботи магістра та має потенціал щодо впровадження результатів у "
    "практичну діяльність підприємств, що працюють із великими обсягами "
    "корпоративної текстової інформації — корпоративні бази знань, системи "
    "технічної підтримки, юридично-довідкові та медичні інформаційні системи.",

    # 15 — substantive contribution
    "Здобувач Нестеренко В.В. провів порівняльне дослідження чотирьох "
    "сучасних моделей векторних ембеддінгів (multilingual-E5-base, BGE-M3, "
    "Qwen3-Embedding-0.6B, nomic-embed-text-v1.5) та лексичної baseline-моделі "
    "BM25 на трьох різнопрофільних доменах знань. Порівняння виконано за "
    "стандартними метриками інформаційного пошуку (nDCG@10, MRR@10, Recall@10, "
    "Precision@10), статистична значущість підтверджена бутстреп-довірчими "
    "інтервалами. Це дозволило обґрунтовано вибрати найбільш придатну модель "
    "для кожного типу корпусу і запропонувати багатокритеріальну процедуру "
    "вибору, що враховує компроміс між якістю пошуку та обчислювальними "
    "витратами.",

    # 16 — practical implementation
    "Проведене дослідження є доцільним і має прикладне значення. Розроблений "
    "веб-застосунок на базі Flask підтверджує можливість практичного "
    "впровадження результатів роботи: він реалізує семантичний пошук, "
    "побудову індексів FAISS, обчислення метрик якості, інтерактивний вибір "
    "моделі за Парето-фронтом і ваговою адитивною моделлю та подетальний "
    "аналіз результатів за окремими запитами. Програма працездатна, придатна "
    "до використання та підтверджує прикладне значення виконаної роботи.",

    # 17 — drawbacks
    "Недоліками роботи можна вважати обмежену кількість релевантних запитів "
    "у бенчмарку за окремими доменами та відсутність експериментів із "
    "наближеними індексами FAISS (HNSW, IVF), що могло б додатково оцінити "
    "поведінку моделей за умов великих корпусів. Однак зазначені обмеження "
    "не применшують наукової та практичної цінності дослідження.",
]

EXTERNAL_CONCLUSION = (
    f"Кваліфікаційна робота магістра, здобувача групи ІПЗм-24-1 "
    f"Нестеренка Віталія Вячеславовича відповідає вимогам до кваліфікаційних "
    f"робіт і заслуговує оцінки {GRADE_PLACEHOLDER}. Кваліфікаційну роботу "
    f"здобувача групи ІПЗм-24-1 Нестеренка В.В. можна представити до захисту "
    f"в ЕК за спеціальністю 121 – «Інженерія програмного забезпечення», "
    f"освітньо-наукова програма «Інженерія програмного забезпечення»."
)

# Two-line external reviewer block. We replace the existing two paragraphs
# (originally degree+name, then chair+univ).
EXTERNAL_REVIEWER_LINE1 = "к.т.н., доцент, \t\t\t\t\t \t\t\tСергій ОРЄХОВ"
EXTERNAL_REVIEWER_LINE2 = (
    "доцент каф. програмної інженерії та інтелектуальних "
    "технологій управління НТУ «ХПІ»"
)


def next_backup(path: pathlib.Path) -> pathlib.Path:
    BACKUPS.mkdir(parents=True, exist_ok=True)
    for n in range(1, 100):
        p = BACKUPS / f"{path.stem}.bak{n}.docx"
        if not p.exists():
            return p
    raise RuntimeError("too many backups")


def set_text(p, text: str, font_pt: float | None = None) -> None:
    """Replace paragraph text, preserve first-run formatting (font size, etc.).

    If font_pt is given, force that point size on the surviving run.
    """
    from docx.shared import Pt
    runs = p.runs
    if runs:
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        runs[0].text = text
        if font_pt is not None:
            runs[0].font.size = Pt(font_pt)
    else:
        run = p.add_run(text)
        if font_pt is not None:
            run.font.size = Pt(font_pt)


def find_idx(doc, predicate) -> int:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    raise RuntimeError("paragraph not found")


def rewrite_internal() -> None:
    backup = next_backup(INT_PATH)
    shutil.copy2(INT_PATH, backup)
    print(f"[backup] {backup.name}")

    d = Document(str(INT_PATH))
    p = d.paragraphs

    # Body of the internal review needs to fit on a single page. The original
    # template uses 14pt; our content is longer than the Romashov original, so
    # we shrink everything below the title to 13pt — matches the visual
    # density of the external review while keeping headings at 14pt.
    BODY_PT = 13.0

    set_text(p[2], GROUP_LINE, font_pt=BODY_PT)
    set_text(p[3], ONP_LINE, font_pt=BODY_PT)
    set_text(p[5], THEME, font_pt=BODY_PT)

    # Static "Представлена робота..." paragraph @ 8 — also shrink for consistency
    set_text(p[8], p[8].text, font_pt=BODY_PT)

    # 4 body paragraphs in slots 9, 10, 11, 12 — no insertions
    for tgt, txt in zip([9, 10, 11, 12], INTERNAL_BODY):
        set_text(p[tgt], txt, font_pt=BODY_PT)

    # Conclusion @ 13
    set_text(p[13], INTERNAL_CONCLUSION, font_pt=BODY_PT)

    # Reviewer @ 19
    rev = find_idx(d, lambda t: "ВЕЧІРСЬКА" in t or "доцент каф. Інформатики" in t)
    set_text(d.paragraphs[rev], INTERNAL_REVIEWER, font_pt=BODY_PT)

    d.save(str(INT_PATH))
    print(f"[saved] {INT_PATH.name}")


def rewrite_external() -> None:
    backup = next_backup(EXT_PATH)
    shutil.copy2(EXT_PATH, backup)
    print(f"[backup] {backup.name}")

    d = Document(str(EXT_PATH))
    p = d.paragraphs

    set_text(p[3], GROUP_LINE)
    set_text(p[4], ONP_LINE)
    set_text(p[7], THEME)

    # 5 body paragraphs 13..17 — 1:1 mapping
    for tgt, txt in zip([13, 14, 15, 16, 17], EXTERNAL_BODY):
        set_text(p[tgt], txt)

    # Conclusion @ 18
    set_text(p[18], EXTERNAL_CONCLUSION)

    # Reviewer block — two paragraphs (name+degree, chair+univ)
    name_idx = find_idx(d, lambda t: "АЛЕКСАНДРОВА" in t)
    set_text(d.paragraphs[name_idx], EXTERNAL_REVIEWER_LINE1)
    chair_idx = find_idx(d, lambda t: "САІТ" in t or "зав.каф." in t)
    set_text(d.paragraphs[chair_idx], EXTERNAL_REVIEWER_LINE2)

    d.save(str(EXT_PATH))
    print(f"[saved] {EXT_PATH.name}")


def main() -> None:
    if not INT_PATH.exists() or not EXT_PATH.exists():
        sys.exit("FAIL: review files not found")
    rewrite_internal()
    rewrite_external()
    print("\nDone.")


if __name__ == "__main__":
    main()
