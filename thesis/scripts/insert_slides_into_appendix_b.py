"""
Insert all 18 presentation slides as images into Додаток Б of the thesis,
two slides per A4 page, preserving 16:9 aspect ratio.

Workflow:
  1. Locate the placeholder paragraph "TODO" that lives just under the
     "Додаток Б — Слайди презентації" heading (around paragraph index 1117).
  2. Replace that TODO with slide #1's image.
  3. After it, append paragraphs with slide images #2 .. #18, applying a
     page break before every odd-numbered slide (3, 5, 7, 9, 11, 13, 15, 17)
     so each A4 page contains exactly two slides.
  4. Save in place. Backups are written as
     thesis/2026_M_PI_Nesterenko_VV.bakNN.docx (next available index).

Image sizing (only width is set; python-docx scales height proportionally):
  page width = 210 mm, left + right margins ≈ 35 mm  ⇒ usable ≈ 175 mm
  we choose 170 mm so the image stops a hair short of the right margin.
  height (16:9) = 170 × 9/16 ≈ 95.6 mm; two images + small gap fits inside
  the 257 mm content height with comfortable breathing room.
"""
import os
import shutil
import sys
import io
import pathlib

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = pathlib.Path("D:/repos/kb-semantic-search-benchmark")
DOCX = ROOT / "thesis" / "2026_M_PI_Nesterenko_VV.docx"
SLIDES_DIR = ROOT / "thesis" / "scripts" / "_slide_imgs"
BACKUPS_DIR = ROOT / "thesis" / "_backups"

IMG_WIDTH_MM = 170  # final width on the page
SLIDES_PER_PAGE = 2


def insert_paragraph_after(paragraph, style=None):
    """Create a new <w:p/> element directly after `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style is not None:
        np.style = style
    return np


def next_backup_path(docx_path: pathlib.Path) -> pathlib.Path:
    """Next free thesis/_backups/<stem>.bakNN.docx slot.

    Backups live in a sibling directory rather than next to the docx itself
    so the working folder stays uncluttered.
    """
    BACKUPS_DIR.mkdir(parents=True, exist_ok=True)
    for n in range(1, 100):
        p = BACKUPS_DIR / f"{docx_path.stem}.bak{n}.docx"
        if not p.exists():
            return p
    raise RuntimeError("Too many backups already.")


def main() -> None:
    if not SLIDES_DIR.exists():
        sys.exit(f"FAIL: slides dir not found: {SLIDES_DIR}")
    slide_paths = [SLIDES_DIR / f"Slide{i}.PNG" for i in range(1, 19)]
    for p in slide_paths:
        if not p.exists():
            sys.exit(f"FAIL: missing slide image {p}")
    print(f"[input] {len(slide_paths)} slide images: {SLIDES_DIR}")

    backup = next_backup_path(DOCX)
    shutil.copy2(DOCX, backup)
    print(f"[backup] -> {backup.name}")

    doc = Document(str(DOCX))
    print(f"[load]  {len(doc.paragraphs)} paragraphs")

    # Locate the TODO placeholder under "Додаток Б". The first occurrence
    # of "Додаток Б" is the TOC entry (single line, ends with a page number),
    # so we want the *last* one — the actual section heading near the end of
    # the document.
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Real heading is multi-line ("Додаток Б\nСлайди презентації") with
        # paragraph style "Додаток"; TOC entry has style ToC and is one line.
        if text.startswith("Додаток Б") and (p.style.name == "Додаток"):
            heading_idx = i
    if heading_idx is None:
        sys.exit("FAIL: 'Додаток Б' heading not found")

    todo_idx = None
    for i in range(heading_idx + 1, min(heading_idx + 10, len(doc.paragraphs))):
        if doc.paragraphs[i].text.strip() == "TODO":
            todo_idx = i
            break
    if todo_idx is None:
        sys.exit("FAIL: 'TODO' placeholder not found under Додаток Б heading")

    todo_para = doc.paragraphs[todo_idx]
    print(f"[anchor] heading @ {heading_idx}, TODO @ {todo_idx}")

    # Clear TODO text content but keep the paragraph element itself (we'll
    # reuse it for slide #1).
    todo_para.clear()
    todo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    prev = todo_para
    img_w = Mm(IMG_WIDTH_MM)
    for i, img_path in enumerate(slide_paths, start=1):
        if i == 1:
            target = prev
        else:
            target = insert_paragraph_after(prev)
            target.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Page break before every odd-indexed slide starting at #3
            if i % SLIDES_PER_PAGE == 1:
                target.paragraph_format.page_break_before = True

        run = target.add_run()
        run.add_picture(str(img_path), width=img_w)
        prev = target
        page_no = (i + 1) // SLIDES_PER_PAGE
        slot = "top" if (i % SLIDES_PER_PAGE) == 1 else "bot"
        print(f"  [{i:2d}] page {page_no} ({slot}): {img_path.name}")

    doc.save(str(DOCX))
    print(f"[save]  {DOCX} ({DOCX.stat().st_size:,} bytes)")
    print("\nDone.")


if __name__ == "__main__":
    main()
