"""
Розширює підсекцію 5.3 «Дослідник benchmark-запитів» описом нової
табної навігації, синхронізації прогонів і drill-down фільтрації за лідером.
Додає 3-й рисунок (5.17) — Дослідник у режимі drill-down BGE-M3.
Перенумеровує існуючі 5.17–5.25 → 5.18–5.26.
"""
import sys, re, pathlib
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

sys.stdout.reconfigure(encoding='utf-8')

ROOT = pathlib.Path("D:/repos/kb-semantic-search-benchmark")
DOCX_PATH = ROOT / "2026_M_PI_Nesterenko_VV.docx"
SHOTS_DIR = ROOT / "docs" / "screenshots"

doc = Document(str(DOCX_PATH))

# ── Step 1: renumber existing 5.17..5.25 → 5.18..5.26 ──────────────
renumber_map = {f"5.{n}": f"5.{n+1}" for n in range(25, 16, -1)}  # reverse order in display
fig_pattern = re.compile(r"5\.(?:17|18|19|20|21|22|23|24|25)")

def replace_in_runs(paragraph):
    changed = False
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = fig_pattern.sub(lambda m: renumber_map.get(m.group(0), m.group(0)), run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    return changed

renumbered = 0
for p in doc.paragraphs:
    if replace_in_runs(p):
        renumbered += 1
print(f"Renumbered references in {renumbered} paragraphs.")

# ── Step 2: find closing paragraph (insertion target) ───────────────
target_para = None
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Реалізований Дослідник benchmark-запитів забезпечує" in p.text:
        target_para = p
        target_idx = i
        break

if target_para is None:
    print("[!] Closing paragraph not found; abort")
    sys.exit(1)

print(f"Found insertion target at paragraph #{target_idx}")
target_xml = target_para._element

# Find reference paragraphs for style
ref_body_para = None
ref_caption_para = None
for p in doc.paragraphs:
    if not ref_body_para and p.text.strip().startswith("Подання результатів"):
        ref_body_para = p
    if not ref_caption_para and "Рисунок 5.14" in p.text:
        ref_caption_para = p
    if ref_body_para and ref_caption_para:
        break

def make_style_para(text: str, ref_para):
    new_p = doc.add_paragraph()
    ref_pPr = ref_para._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_pPr = new_p._element.find(qn("w:pPr"))
        if new_pPr is not None:
            new_p._element.remove(new_pPr)
        new_p._element.insert(0, deepcopy(ref_pPr))
    if text:
        run = new_p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    return new_p

# ── Step 3: build new content ───────────────────────────────────────
intro_text = (
    "Дослідник benchmark-запитів інтегровано з основною сторінкою Benchmark через "
    "спільну табну навігацію — вкладки «Огляд моделей» і «По запитах» (рисунки 5.15–5.16). "
    "Така організація дозволяє користувачу швидко перемикатися між агрегованим переглядом "
    "моделей і деталізованим аналізом запитів того самого прогону без втрати контексту. "
    "У правій частині табної смужки реалізовано випадаючий список історичних benchmark-"
    "прогонів: обраний прогін зберігається при перемиканні між вкладками, що забезпечує "
    "консистентність аналізу при порівнянні запусків експерименту в часі."
)
intro_para = make_style_para(intro_text, ref_body_para)

drill_text = (
    "Окрему функціональну можливість становить drill-down навігація: біля кожної моделі "
    "у зведеній таблиці Огляду моделей розміщено піктограму переходу, активація якої "
    "відкриває Дослідник із автоматично застосованим фільтром «Лідер: <модель>». "
    "Результат фільтрації наведено на рисунку 5.17 на прикладі моделі BGE-M3: показано "
    "лише ті benchmark-запити, на яких ця модель випередила всі інші за метрикою nDCG@10. "
    "Над таблицею відображено інформаційний банер з лічильником кількості таких запитів, "
    "що дозволяє швидко оцінити частку benchmark-набору, на якій кожна модель є "
    "найкращою. Аналогічна інформація доступна у вигляді «пігулок-фільтрів» з "
    "лічильниками у рядку «Модель» над таблицею запитів — це робить розподіл лідерства "
    "за моделями наочним з першого погляду."
)
drill_para = make_style_para(drill_text, ref_body_para)

# Image 3 (drill-down)
img3_para = doc.add_paragraph()
img3_run = img3_para.add_run()
img3_run.add_picture(str(SHOTS_DIR / "explorer_drill_bge_m3.png"), width=Inches(6.0))
img3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

caption3_text = (
    "Рисунок 5.17 – Дослідник у режимі drill-down фільтрації за лідером: запити "
    "тестового набору, на яких модель BGE-M3 випередила інші за nDCG@10 (рисунок "
    "створено самостійно)"
)
caption3_para = make_style_para(caption3_text, ref_caption_para)

# ── Step 4: insert before closing paragraph ─────────────────────────
new_paragraphs = [intro_para, drill_para, img3_para, caption3_para]

for new_p in new_paragraphs:
    new_xml = new_p._element
    new_xml.getparent().remove(new_xml)
    target_xml.addprevious(new_xml)

print(f"Inserted {len(new_paragraphs)} new paragraphs before closing.")

# ── Step 5: save ──────────────────────────────────────────────────
doc.save(str(DOCX_PATH))
print(f"Saved: {DOCX_PATH} ({DOCX_PATH.stat().st_size:,} bytes)")
print("DONE.")
