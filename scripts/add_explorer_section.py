"""
Розширює розділ 5.3 «Підхід до розробки» новою підсекцією про
«Дослідник benchmark-запитів» (Benchmark Explorer) з 2 скріншотами.

  1. Вставляє перед підсумковим абзацом (B988 «Отже, реалізований підхід…»)
     блок з 2 пояснювальних абзаців + 2 рисунки (5.15, 5.16).
  2. Усі існуючі рисунки 5.15–5.23 перенумеровуються у 5.17–5.25
     (як у підписах, так і у внутрішньотекстових посиланнях).
"""
import sys, re, copy, pathlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

ROOT = pathlib.Path("D:/repos/kb-semantic-search-benchmark")
DOCX_PATH = ROOT / "2026_M_PI_Nesterenko_VV.docx"
SHOTS_DIR = ROOT / "docs" / "screenshots"

# ── Step 1: open document and locate insertion point ─────────────────────
doc = Document(str(DOCX_PATH))
body = doc.element.body

# Find paragraph that starts with "Отже, реалізований підхід до розробки"
target_para = None
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Отже, реалізований підхід до розробки" in p.text:
        target_para = p
        target_idx = i
        break

if target_para is None:
    print("[!] Target paragraph not found; abort")
    sys.exit(1)

print(f"Found insertion target at paragraph #{target_idx}")
target_xml = target_para._element  # lxml element

# ── Step 2: renumber existing figures 5.15+ → +2 ───────────────────────
# Map old → new
renumber_map = {
    "5.23": "5.25",
    "5.22": "5.24",
    "5.21": "5.23",
    "5.20": "5.22",
    "5.19": "5.21",
    "5.18": "5.20",
    "5.17": "5.19",
    "5.16": "5.18",
    "5.15": "5.17",
}

# Apply via regex on every run text in document
fig_pattern = re.compile(r"5\.(?:15|16|17|18|19|20|21|22|23)")

def replace_in_runs(paragraph):
    """Renumber figure references in a paragraph's runs."""
    changed = False
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = fig_pattern.sub(lambda m: renumber_map.get(m.group(0), m.group(0)), run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    return changed

renumbered = 0
for p in doc.paragraphs:
    if replace_in_runs(p):
        renumbered += 1
print(f"Renumbered references in {renumbered} paragraphs.")

# ── Step 3: build the new content (2 paragraphs + 2 figures + 2 captions) ──

# Helper: copy paragraph style from a reference paragraph
def make_style_para(text: str, ref_para):
    """Create a new paragraph with same style as ref_para, containing text."""
    # Use python-docx API: add a paragraph at end, then move it
    new_p = doc.add_paragraph()
    # copy paragraph properties from ref
    ref_pPr = ref_para._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_pPr = new_p._element.find(qn("w:pPr"))
        if new_pPr is not None:
            new_p._element.remove(new_pPr)
        new_p._element.insert(0, deepcopy(ref_pPr))
    # add run with text
    if text:
        run = new_p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    return new_p


# Find a good reference paragraph (regular body) and a caption paragraph for style
ref_body_para = None
ref_caption_para = None
for p in doc.paragraphs:
    if not ref_body_para and p.text.strip().startswith("Подання результатів"):
        ref_body_para = p
    if not ref_caption_para and "Рисунок 5.14" in p.text:
        ref_caption_para = p
    if ref_body_para and ref_caption_para:
        break

# ── New paragraphs ───────────────────────────────────────────────────────
intro_text = (
    "Окрім зведених benchmark-метрик, у системі реалізовано інструмент детального "
    "перегляду тестових даних і фактичних результатів пошуку — Дослідник benchmark-"
    "запитів. Він дозволяє переходити від агрегованих числових показників до "
    "конкретних запитів, еталонних релевантних відповідностей (qrels) і top-k "
    "результатів кожної моделі. Це робить benchmark-оцінювання прозорим: "
    "користувач може перевірити, який саме чанк документа модель повернула на "
    "перше місце, чи збігається він з еталоном, і як інші моделі впоралися із "
    "тим самим запитом."
)
intro_para = make_style_para(intro_text, ref_body_para)

before_caption_text = (
    "Список усіх benchmark-запитів обраного домену з підсумковими per-query "
    "метриками наведено на рисунку 5.15. На сторінці відображено унікальні ID "
    "запитів, текст, кількість релевантних чанків (qrels), значення nDCG@10 "
    "найкращої моделі, розкид метрик між моделями та лідера. Доступні фільтри "
    "за категоріями (ідеальні, провальні, з великим розкидом моделей) і "
    "сортування за різними критеріями."
)
list_para = make_style_para(before_caption_text, ref_body_para)

# Image 1 paragraph (centered, image, no text)
img1_para = doc.add_paragraph()
img1_run = img1_para.add_run()
img1_run.add_picture(str(SHOTS_DIR / "explorer_index_tech.png"), width=Inches(6.0))
img1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Caption 1
caption1_text = (
    "Рисунок 5.15 – Список benchmark-запитів у Досліднику з per-query метриками "
    "та фільтрами (рисунок створено самостійно)"
)
caption1_para = make_style_para(caption1_text, ref_caption_para)

# Detail paragraph
detail_text = (
    "Деталізована сторінка одного benchmark-запиту наведена на рисунку 5.16. На ній "
    "показано текст запиту, ground-truth релевантних чанків і вкладки з результатами "
    "кожної моделі. У вкладці моделі відображаються: per-query метрики (nDCG@10, "
    "MRR@10, Recall@10, P@10, час відповіді) і top-10 знайдених чанків. Чанки, які "
    "збігаються з qrels, виділяються зеленим, що дозволяє візуально оцінити, чи "
    "модель повернула справді релевантний контент і на якій позиції."
)
detail_para = make_style_para(detail_text, ref_body_para)

# Image 2
img2_para = doc.add_paragraph()
img2_run = img2_para.add_run()
img2_run.add_picture(str(SHOTS_DIR / "explorer_query_q1.png"), width=Inches(6.0))
img2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Caption 2
caption2_text = (
    "Рисунок 5.16 – Деталізована сторінка benchmark-запиту з top-10 результатами "
    "кожної моделі та підсвіткою релевантних чанків (рисунок створено самостійно)"
)
caption2_para = make_style_para(caption2_text, ref_caption_para)

# Closing context paragraph
closing_text = (
    "Реалізований Дослідник benchmark-запитів забезпечує методичну прозорість "
    "експерименту: будь-яке агреговане значення метрики можна простежити до конкретних "
    "запитів і фактично знайдених документів. Цей інструмент є ключовим для якісного "
    "аналізу помилок моделей і обґрунтування підсумкових висновків дослідження."
)
closing_para = make_style_para(closing_text, ref_body_para)

# ── Step 4: move all the new paragraphs BEFORE target_xml ───────────────
new_paragraphs = [
    intro_para, list_para, img1_para, caption1_para,
    detail_para, img2_para, caption2_para, closing_para,
]

for new_p in new_paragraphs:
    new_xml = new_p._element
    # Remove from current position (was appended at end)
    new_xml.getparent().remove(new_xml)
    # Insert before target
    target_xml.addprevious(new_xml)

print(f"Inserted {len(new_paragraphs)} new paragraphs before target.")

# ── Step 5: save ────────────────────────────────────────────────────────
doc.save(str(DOCX_PATH))
print(f"Saved: {DOCX_PATH} ({DOCX_PATH.stat().st_size:,} bytes)")
print("DONE.")
